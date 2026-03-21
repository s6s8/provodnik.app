import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
CONTENT_PATH = ROOT / "docs" / "presentations" / "provodnik-proekt-i-plan-mvp-ru.json"
OUTPUT_DOCX = ROOT / "docs" / "presentations" / "provodnik-proekt-i-plan-mvp-ru.docx"
OUTPUT_PDF = ROOT / "docs" / "presentations" / "provodnik-proekt-i-plan-mvp-ru.pdf"

NAVY = RGBColor(0x1D, 0x34, 0x47)
TEAL = RGBColor(0x0F, 0x76, 0x6E)
GRAY = RGBColor(0x5C, 0x67, 0x70)


def load_content() -> dict:
    return json.loads(CONTENT_PATH.read_text(encoding="utf-8"))


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def configure_docx_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)

    style_map = [
        ("Heading 1", 17, NAVY),
        ("Heading 2", 13, NAVY),
        ("Heading 3", 11, TEAL),
    ]
    for style_name, size, color in style_map:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    if "Small Text" not in styles:
        small = styles.add_style("Small Text", WD_STYLE_TYPE.PARAGRAPH)
        small.base_style = styles["Normal"]
        small.font.name = "Arial"
        small._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        small.font.size = Pt(9)
        small.font.color.rgb = GRAY


def add_bullet(document: Document, text: str, level: int = 0) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run(text)


def build_docx(content: dict) -> None:
    document = Document()
    for section in document.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    configure_docx_styles(document)
    meta = content["meta"]

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(meta["title"])
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = NAVY

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(meta["subtitle"])
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(13)
    run.font.color.rgb = GRAY

    tagline = document.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run(meta["tagline"])
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(11)
    run.font.color.rgb = TEAL

    document.add_paragraph("")
    document.add_paragraph("Краткое резюме", style="Heading 1")
    for item in content["summary"]:
        add_bullet(document, item)

    document.add_page_break()

    for chapter in content["chapters"]:
        document.add_paragraph(chapter["title"], style="Heading 1")
        for paragraph_text in chapter.get("paragraphs", []):
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.add_run(paragraph_text)
        for bullet in chapter.get("bullets", []):
            add_bullet(document, bullet)

        for subsection in chapter.get("subsections", []):
            document.add_paragraph(subsection["title"], style="Heading 2")
            for item in subsection["items"]:
                add_bullet(document, item, level=1)

        for stage in chapter.get("stages", []):
            document.add_paragraph(stage["name"], style="Heading 2")
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run("Цель: ")
            run.bold = True
            p.add_run(stage["goal"])
            for step in stage["steps"]:
                add_bullet(document, step, level=1)

        document.add_paragraph("")

    document.add_paragraph(content["appendix"]["title"], style="Heading 1")
    for item in content["appendix"]["items"]:
        add_bullet(document, item)

    document.add_paragraph("Источники", style="Heading 1")
    for source in content["sources"]:
        add_bullet(document, source)

    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.style = document.styles["Small Text"]
        footer.add_run("Страница ")
        add_page_number(footer)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)


def find_font(*candidates: str) -> str:
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    raise FileNotFoundError(f"Missing font from candidates: {candidates}")


def register_fonts() -> None:
    regular = find_font(
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\segoeui.ttf",
        r"C:\Windows\Fonts\tahoma.ttf",
    )
    bold = find_font(
        r"C:\Windows\Fonts\arialbd.ttf",
        r"C:\Windows\Fonts\segoeuib.ttf",
        r"C:\Windows\Fonts\tahomabd.ttf",
    )
    pdfmetrics.registerFont(TTFont("ProjRegular", regular))
    pdfmetrics.registerFont(TTFont("ProjBold", bold))


def draw_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor("#DAD4CA"))
    canvas.setLineWidth(0.6)
    canvas.line(18 * mm, 15 * mm, 192 * mm, 15 * mm)
    canvas.setFont("ProjRegular", 8)
    canvas.setFillColor(colors.HexColor("#5C6770"))
    canvas.drawString(18 * mm, 10.5 * mm, "Проводник | Проект и план минимально жизнеспособной версии")
    canvas.drawRightString(192 * mm, 10.5 * mm, f"Страница {doc.page}")
    canvas.restoreState()


def build_pdf(content: dict) -> None:
    register_fonts()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontName="ProjBold",
        fontSize=24,
        leading=28,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#1D3447"),
        spaceAfter=10,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["BodyText"],
        fontName="ProjRegular",
        fontSize=12,
        leading=15,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#5C6770"),
        spaceAfter=5,
    )
    accent_style = ParagraphStyle(
        "Accent",
        parent=styles["BodyText"],
        fontName="ProjBold",
        fontSize=10,
        leading=12,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#0F766E"),
        spaceAfter=16,
    )
    h1 = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontName="ProjBold",
        fontSize=16,
        leading=19,
        textColor=colors.HexColor("#1D3447"),
        spaceAfter=8,
        spaceBefore=4,
    )
    h2 = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontName="ProjBold",
        fontSize=12,
        leading=15,
        textColor=colors.HexColor("#0F766E"),
        spaceAfter=6,
        spaceBefore=4,
    )
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="ProjRegular",
        fontSize=10,
        leading=13.5,
        textColor=colors.HexColor("#13212B"),
        spaceAfter=7,
    )
    bullet = ParagraphStyle(
        "Bullet",
        parent=body,
        leftIndent=13,
        bulletIndent=1,
        spaceAfter=3,
    )
    small = ParagraphStyle(
        "Small",
        parent=body,
        fontSize=9,
        leading=11.5,
        textColor=colors.HexColor("#5C6770"),
    )

    doc = SimpleDocTemplate(
        str(OUTPUT_PDF),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=20 * mm,
    )

    story = []
    meta = content["meta"]
    story.append(Spacer(1, 32))
    story.append(Paragraph(meta["title"], title_style))
    story.append(Paragraph(meta["subtitle"], subtitle_style))
    story.append(Paragraph(meta["tagline"], accent_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Краткое резюме", h1))
    for item in content["summary"]:
        story.append(Paragraph(item, bullet, bulletText="-"))
    story.append(PageBreak())

    for chapter in content["chapters"]:
        story.append(Paragraph(chapter["title"], h1))
        for paragraph_text in chapter.get("paragraphs", []):
            story.append(Paragraph(paragraph_text, body))
        for item in chapter.get("bullets", []):
            story.append(Paragraph(item, bullet, bulletText="-"))

        for subsection in chapter.get("subsections", []):
            story.append(Paragraph(subsection["title"], h2))
            for item in subsection["items"]:
                story.append(Paragraph(item, bullet, bulletText="-"))

        for stage in chapter.get("stages", []):
            story.append(Paragraph(stage["name"], h2))
            story.append(Paragraph(f"<b>Цель:</b> {stage['goal']}", body))
            for step in stage["steps"]:
                story.append(Paragraph(step, bullet, bulletText="-"))

        story.append(Spacer(1, 7))

    story.append(Paragraph(content["appendix"]["title"], h1))
    for item in content["appendix"]["items"]:
        story.append(Paragraph(item, bullet, bulletText="-"))

    story.append(Paragraph("Источники", h1))
    for source in content["sources"]:
        story.append(Paragraph(source, bullet, bulletText="-"))
    story.append(
        Paragraph(
            f"Документ собран на основе внутренних материалов проекта и research snapshot от {meta['basis_date']}.",
            small,
        )
    )

    OUTPUT_PDF.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)


def main() -> None:
    content = load_content()
    build_docx(content)
    build_pdf(content)
    print(f"Generated {OUTPUT_DOCX}")
    print(f"Generated {OUTPUT_PDF}")


if __name__ == "__main__":
    main()
