import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.utils import simpleSplit
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
CONTENT_PATH = ROOT / "docs" / "presentations" / "provodnik-marketing-research-ru.json"
OUTPUT_DOCX = ROOT / "docs" / "presentations" / "provodnik-marketing-research-ru.docx"
OUTPUT_PDF = ROOT / "docs" / "presentations" / "provodnik-marketing-research-ru.pdf"

COLORS = {
    "charcoal": RGBColor(0x20, 0x26, 0x2D),
    "sand": RGBColor(0xF2, 0xED, 0xE4),
    "brick": RGBColor(0xB9, 0x5C, 0x47),
    "olive": RGBColor(0x6E, 0x7A, 0x52),
    "slate": RGBColor(0x4C, 0x58, 0x66),
}

PDF_COLORS = {
    "charcoal": HexColor("#20262D"),
    "sand": HexColor("#F2EDE4"),
    "brick": HexColor("#B95C47"),
    "olive": HexColor("#6E7A52"),
    "slate": HexColor("#4C5866"),
    "paper": HexColor("#FAF7F2"),
    "line": HexColor("#D8CFC2"),
    "white": HexColor("#FFFFFF"),
}


def load_content() -> dict:
    return json.loads(CONTENT_PATH.read_text(encoding="utf-8"))


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)

    for style_name, size, color in [
        ("Heading 1", 22, COLORS["charcoal"]),
        ("Heading 2", 14, COLORS["brick"]),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    if "Slide Small" not in styles:
        small = styles.add_style("Slide Small", WD_STYLE_TYPE.PARAGRAPH)
        small.base_style = styles["Normal"]
        small.font.name = "Arial"
        small._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        small.font.size = Pt(9)
        small.font.color.rgb = COLORS["slate"]


def cover_page_docx(document: Document, content: dict) -> None:
    meta = content["meta"]
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta["title"])
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLORS["charcoal"]

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta["subtitle"])
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(15)
    run.font.color.rgb = COLORS["slate"]

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta["tagline"])
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(11)
    run.font.color.rgb = COLORS["brick"]

    document.add_paragraph("")
    document.add_paragraph("Повестка", style="Heading 2")
    table = document.add_table(rows=2, cols=4)
    table.autofit = False
    widths = [Cm(6.3), Cm(6.3), Cm(6.3), Cm(6.3)]
    idx = 0
    for row in table.rows:
        row.height = Cm(2.1)
        for cell, width in zip(row.cells, widths):
            cell.width = width
            if idx < len(content["agenda"]):
                cell.text = content["agenda"][idx]
                set_cell_shading(cell, "F2EDE4")
            idx += 1

    document.add_paragraph("")
    document.add_paragraph("Основа исследования", style="Heading 2")
    for item in content["methodology"]:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_cards_page_docx(document: Document, slide: dict) -> None:
    document.add_paragraph(slide["title"], style="Heading 1")
    document.add_paragraph(slide["lead"])
    table = document.add_table(rows=2, cols=2)
    table.autofit = False
    widths = [Cm(12.8), Cm(12.8)]
    cards = slide["cards"]
    idx = 0
    for row in table.rows:
        row.height = Cm(4.0)
        for cell, width in zip(row.cells, widths):
            cell.width = width
            if idx < len(cards):
                card = cards[idx]
                set_cell_shading(cell, "F7F3EC")
                cell.paragraphs[0].add_run(card["label"] + "\n").bold = True
                title_run = cell.paragraphs[0].add_run(card["title"] + "\n")
                title_run.bold = True
                cell.paragraphs[0].add_run(card["body"])
            idx += 1


def add_split_page_docx(document: Document, slide: dict) -> None:
    document.add_paragraph(slide["title"], style="Heading 1")
    document.add_paragraph(slide["lead"])
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    row = table.rows[0]
    row.height = Cm(8.5)
    for cell in row.cells:
        cell.width = Cm(12.8)
        set_cell_shading(cell, "F7F3EC")

    left = row.cells[0].paragraphs[0]
    left.add_run(slide["left_title"] + "\n").bold = True
    for item in slide["left_points"]:
        left.add_run("• " + item + "\n")

    right = row.cells[1].paragraphs[0]
    right.add_run(slide["right_title"] + "\n").bold = True
    for item in slide["right_points"]:
        right.add_run("• " + item + "\n")


def build_docx(content: dict) -> None:
    document = Document()
    configure_docx(document)
    cover_page_docx(document, content)

    for slide in content["slides"]:
        document.add_page_break()
        if "cards" in slide:
            add_cards_page_docx(document, slide)
        else:
            add_split_page_docx(document, slide)

    document.add_page_break()
    document.add_paragraph("Источники и примечание", style="Heading 1")
    document.add_paragraph(content["sources_note"])
    for item in content["sources_note"].split(". "):
        if item.strip():
            p = document.add_paragraph(style="List Bullet")
            p.add_run(item.strip().rstrip(".") + ".")

    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.style = document.styles["Slide Small"]
        footer.add_run("Страница ")
        add_page_number(footer)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)


def register_fonts() -> None:
    regular = Path(r"C:\Windows\Fonts\arial.ttf")
    bold = Path(r"C:\Windows\Fonts\arialbd.ttf")
    if not regular.exists() or not bold.exists():
        raise FileNotFoundError("Arial fonts not found.")
    pdfmetrics.registerFont(TTFont("DeckRegular", str(regular)))
    pdfmetrics.registerFont(TTFont("DeckBold", str(bold)))


def wrap(text: str, font: str, size: int, width: float) -> list[str]:
    return simpleSplit(text, font, size, width)


def draw_text_block(c: canvas.Canvas, x: float, y: float, width: float, lines: list[str], font: str, size: int, leading: float, color) -> float:
    c.setFont(font, size)
    c.setFillColor(color)
    current_y = y
    for line in lines:
        c.drawString(x, current_y, line)
        current_y -= leading
    return current_y


def draw_cover_pdf(c: canvas.Canvas, width: float, height: float, content: dict) -> None:
    meta = content["meta"]
    c.setFillColor(PDF_COLORS["charcoal"])
    c.rect(0, 0, width, height, fill=1, stroke=0)
    c.setFillColor(PDF_COLORS["sand"])
    c.rect(18, 18, width - 36, height - 36, fill=0, stroke=1)
    c.setFillColor(PDF_COLORS["white"])
    c.setFont("DeckBold", 28)
    c.drawString(42, height - 70, meta["title"])
    c.setFont("DeckBold", 18)
    for idx, line in enumerate(wrap(meta["subtitle"], "DeckBold", 18, width - 90)):
        c.drawString(42, height - 105 - idx * 23, line)
    c.setFont("DeckRegular", 11)
    c.setFillColor(PDF_COLORS["sand"])
    c.drawString(42, height - 160, meta["tagline"])
    c.setFillColor(PDF_COLORS["brick"])
    c.roundRect(42, height - 330, width - 84, 126, 16, fill=0, stroke=1)
    c.setFont("DeckBold", 13)
    c.drawString(58, height - 225, "Повестка")
    c.setFont("DeckRegular", 11)
    agenda = content["agenda"]
    for i, item in enumerate(agenda):
        col = i % 2
        row = i // 2
        x = 58 + col * 240
        y = height - 250 - row * 22
        c.drawString(x, y, f"{i+1}. {item}")
    c.setFont("DeckBold", 13)
    c.drawString(58, 138, "Основа исследования")
    c.setFont("DeckRegular", 10)
    y = 118
    for item in content["methodology"]:
        lines = wrap("• " + item, "DeckRegular", 10, width - 120)
        y = draw_text_block(c, 58, y, width - 120, lines, "DeckRegular", 10, 13, PDF_COLORS["sand"]) - 5


def draw_slide_header(c: canvas.Canvas, width: float, height: float, title: str, lead: str) -> float:
    c.setFillColor(PDF_COLORS["paper"])
    c.rect(0, 0, width, height, fill=1, stroke=0)
    c.setFillColor(PDF_COLORS["charcoal"])
    c.rect(0, height - 66, width, 66, fill=1, stroke=0)
    c.setFillColor(PDF_COLORS["white"])
    c.setFont("DeckBold", 20)
    title_lines = wrap(title, "DeckBold", 20, width - 80)
    yy = height - 42
    for line in title_lines:
        c.drawString(36, yy, line)
        yy -= 22
    c.setFillColor(PDF_COLORS["slate"])
    c.setFont("DeckRegular", 11)
    lead_lines = wrap(lead, "DeckRegular", 11, width - 84)
    y = height - 94
    for line in lead_lines:
        c.drawString(36, y, line)
        y -= 14
    return y - 8


def draw_cards_slide(c: canvas.Canvas, width: float, height: float, slide: dict) -> None:
    y = draw_slide_header(c, width, height, slide["title"], slide["lead"])
    card_w = (width - 36 * 2 - 16) / 2
    card_h = 120
    positions = [
        (36, y - card_h),
        (36 + card_w + 16, y - card_h),
        (36, y - card_h * 2 - 14),
        (36 + card_w + 16, y - card_h * 2 - 14),
    ]
    for card, (x, yy) in zip(slide["cards"], positions):
        c.setFillColor(PDF_COLORS["white"])
        c.setStrokeColor(PDF_COLORS["line"])
        c.roundRect(x, yy, card_w, card_h, 14, fill=1, stroke=1)
        c.setFillColor(PDF_COLORS["brick"])
        c.setFont("DeckBold", 9)
        c.drawString(x + 14, yy + card_h - 18, card["label"].upper())
        c.setFillColor(PDF_COLORS["charcoal"])
        c.setFont("DeckBold", 14)
        title_lines = wrap(card["title"], "DeckBold", 14, card_w - 28)
        text_y = yy + card_h - 40
        for line in title_lines:
            c.drawString(x + 14, text_y, line)
            text_y -= 16
        body_lines = wrap(card["body"], "DeckRegular", 10, card_w - 28)
        c.setFont("DeckRegular", 10)
        c.setFillColor(PDF_COLORS["slate"])
        for line in body_lines:
            c.drawString(x + 14, text_y, line)
            text_y -= 13


def draw_split_slide(c: canvas.Canvas, width: float, height: float, slide: dict) -> None:
    y = draw_slide_header(c, width, height, slide["title"], slide["lead"])
    box_w = (width - 36 * 2 - 16) / 2
    box_h = 245
    for idx, key in enumerate(["left", "right"]):
        x = 36 + idx * (box_w + 16)
        c.setFillColor(PDF_COLORS["white"])
        c.setStrokeColor(PDF_COLORS["line"])
        c.roundRect(x, y - box_h, box_w, box_h, 14, fill=1, stroke=1)
        c.setFillColor(PDF_COLORS["olive"] if key == "left" else PDF_COLORS["brick"])
        c.setFont("DeckBold", 13)
        c.drawString(x + 14, y - 24, slide[f"{key}_title"])
        current_y = y - 46
        c.setFont("DeckRegular", 10)
        c.setFillColor(PDF_COLORS["charcoal"])
        for item in slide[f"{key}_points"]:
            lines = wrap("• " + item, "DeckRegular", 10, box_w - 28)
            for line in lines:
                c.drawString(x + 14, current_y, line)
                current_y -= 13
            current_y -= 6


def draw_pdf_footer(c: canvas.Canvas, width: float, height: float, page_num: int) -> None:
    c.setStrokeColor(PDF_COLORS["line"])
    c.line(36, 24, width - 36, 24)
    c.setFont("DeckRegular", 8)
    c.setFillColor(PDF_COLORS["slate"])
    c.drawString(36, 12, "Проводник | Маркетинговое исследование")
    c.drawRightString(width - 36, 12, f"Страница {page_num}")


def build_pdf(content: dict) -> None:
    register_fonts()
    width, height = landscape(A4)
    c = canvas.Canvas(str(OUTPUT_PDF), pagesize=(width, height))
    page_num = 1
    draw_cover_pdf(c, width, height, content)
    draw_pdf_footer(c, width, height, page_num)
    c.showPage()
    page_num += 1

    for slide in content["slides"]:
        if "cards" in slide:
            draw_cards_slide(c, width, height, slide)
        else:
            draw_split_slide(c, width, height, slide)
        draw_pdf_footer(c, width, height, page_num)
        c.showPage()
        page_num += 1

    c.save()


def main() -> None:
    content = load_content()
    build_docx(content)
    build_pdf(content)
    print(f"Generated {OUTPUT_DOCX}")
    print(f"Generated {OUTPUT_PDF}")


if __name__ == "__main__":
    main()
