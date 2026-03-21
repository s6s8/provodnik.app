import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
CONTENT_PATH = ROOT / "docs" / "presentations" / "provodnik-project-description-ru.json"
OUTPUT_DOCX = ROOT / "docs" / "presentations" / "provodnik-project-description-ru.docx"
OUTPUT_PDF = ROOT / "docs" / "presentations" / "provodnik-project-description-ru.pdf"

NAVY = RGBColor(0x1D, 0x34, 0x47)
TEAL = RGBColor(0x0F, 0x76, 0x6E)
GRAY = RGBColor(0x5C, 0x67, 0x70)


def load_content() -> dict:
    return json.loads(CONTENT_PATH.read_text(encoding="utf-8"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)


def configure_docx_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Heading 1", 18, NAVY),
        ("Heading 2", 14, NAVY),
        ("Heading 3", 11, TEAL),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    if "Body Small" not in styles:
        small = styles.add_style("Body Small", WD_STYLE_TYPE.PARAGRAPH)
        small.base_style = styles["Normal"]
        small.font.name = "Arial"
        small._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        small.font.size = Pt(9)
        small.font.color.rgb = GRAY


def build_docx(content: dict) -> None:
    document = Document()
    for section in document.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    configure_docx_styles(document)

    meta = content["meta"]
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(meta["title"])
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(meta["subtitle"])
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(13)
    run.font.color.rgb = GRAY

    tagline = document.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run(meta["tagline"])
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(11)
    run.font.color.rgb = TEAL

    document.add_paragraph("")

    intro = document.add_paragraph(style="Heading 2")
    intro.add_run("Краткое резюме")
    for bullet in content["executive_summary"]:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(bullet)

    document.add_page_break()

    for index, section_data in enumerate(content["sections"], start=1):
        document.add_paragraph(section_data["title"], style="Heading 1")
        for paragraph in section_data["paragraphs"]:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.add_run(paragraph)
        for bullet in section_data["bullets"]:
            p = document.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            p.add_run(bullet)

        if index == 3:
            table_data = content["tables"]["competitors"]
            document.add_paragraph(table_data["title"], style="Heading 2")
            table = document.add_table(rows=1, cols=len(table_data["headers"]))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            header_cells = table.rows[0].cells
            for i, header in enumerate(table_data["headers"]):
                header_cells[i].text = header
                set_cell_shading(header_cells[i], "DCEEE8")
            for row in table_data["rows"]:
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    cells[i].text = value

        if index == 6:
            table_data = content["tables"]["roadmap"]
            document.add_paragraph(table_data["title"], style="Heading 2")
            table = document.add_table(rows=1, cols=len(table_data["headers"]))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            header_cells = table.rows[0].cells
            for i, header in enumerate(table_data["headers"]):
                header_cells[i].text = header
                set_cell_shading(header_cells[i], "E9EEF3")
            for row in table_data["rows"]:
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    cells[i].text = value

        if index != len(content["sections"]):
            document.add_paragraph("")

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Источники", style="Heading 1")
    for source in content["sources"]:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(source)

    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.style = document.styles["Body Small"]
        footer.add_run("Страница ")
        add_page_number(footer)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)


def find_font(*candidates: str) -> str:
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    raise FileNotFoundError(f"Missing font from candidates: {candidates}")


def register_fonts() -> None:
    regular = find_font(
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\segoeui.ttf",
        r"C:\Windows\Fonts\tahoma.ttf",
    )
    bold = find_font(
        r"C:\Windows\Fonts\arialbd.ttf",
        r"C:\Windows\Fonts\segoeuib.ttf",
        r"C:\Windows\Fonts\tahomabd.ttf",
    )
    pdfmetrics.registerFont(TTFont("ProvodnikRegular", regular))
    pdfmetrics.registerFont(TTFont("ProvodnikBold", bold))


def pdf_header_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor("#DAD4CA"))
    canvas.setLineWidth(0.6)
    canvas.line(18 * mm, 15 * mm, 192 * mm, 15 * mm)
    canvas.setFont("ProvodnikRegular", 8)
    canvas.setFillColor(colors.HexColor("#5C6770"))
    canvas.drawString(18 * mm, 10.5 * mm, "Provodnik | Описание проекта, продукта и рынка")
    canvas.drawRightString(192 * mm, 10.5 * mm, f"Страница {doc.page}")
    canvas.restoreState()


def make_table(data: dict, col_widths: list[float]) -> Table:
    rows = [data["headers"], *data["rows"]]
    table = Table(rows, colWidths=col_widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DCEEE8")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1D3447")),
                ("FONTNAME", (0, 0), (-1, 0), "ProvodnikBold"),
                ("FONTNAME", (0, 1), (-1, -1), "ProvodnikRegular"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.2),
                ("LEADING", (0, 0), (-1, -1), 10),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#DAD4CA")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def build_pdf(content: dict) -> None:
    register_fonts()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleStyle",
        parent=styles["Title"],
        fontName="ProvodnikBold",
        fontSize=24,
        leading=28,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#1D3447"),
        spaceAfter=12,
    )
    subtitle_style = ParagraphStyle(
        "SubtitleStyle",
        parent=styles["BodyText"],
        fontName="ProvodnikRegular",
        fontSize=12,
        leading=15,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#5C6770"),
        spaceAfter=6,
    )
    tag_style = ParagraphStyle(
        "TagStyle",
        parent=styles["BodyText"],
        fontName="ProvodnikBold",
        fontSize=10,
        leading=12,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#0F766E"),
        spaceAfter=16,
    )
    h1 = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontName="ProvodnikBold",
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#1D3447"),
        spaceAfter=8,
        spaceBefore=4,
    )
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="ProvodnikRegular",
        fontSize=10,
        leading=13.5,
        textColor=colors.HexColor("#13212B"),
        spaceAfter=7,
    )
    bullet = ParagraphStyle(
        "Bullet",
        parent=body,
        leftIndent=14,
        bulletIndent=2,
        spaceAfter=3,
    )
    small = ParagraphStyle(
        "Small",
        parent=body,
        fontSize=9,
        leading=11.5,
        textColor=colors.HexColor("#5C6770"),
    )

    doc = SimpleDocTemplate(
        str(OUTPUT_PDF),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=20 * mm,
    )

    story = []
    meta = content["meta"]
    story.append(Spacer(1, 32))
    story.append(Paragraph(meta["title"], title_style))
    story.append(Paragraph(meta["subtitle"], subtitle_style))
    story.append(Paragraph(meta["tagline"], tag_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Краткое резюме", h1))
    for item in content["executive_summary"]:
        story.append(Paragraph(item, bullet, bulletText="•"))
    story.append(PageBreak())

    for index, section_data in enumerate(content["sections"], start=1):
        story.append(Paragraph(section_data["title"], h1))
        for paragraph in section_data["paragraphs"]:
            story.append(Paragraph(paragraph, body))
        for item in section_data["bullets"]:
            story.append(Paragraph(item, bullet, bulletText="•"))

        if index == 3:
            story.append(Spacer(1, 6))
            story.append(Paragraph(content["tables"]["competitors"]["title"], h1))
            story.append(
                make_table(
                    content["tables"]["competitors"],
                    [38 * mm, 43 * mm, 43 * mm, 43 * mm],
                )
            )
        if index == 6:
            story.append(Spacer(1, 6))
            story.append(Paragraph(content["tables"]["roadmap"]["title"], h1))
            story.append(
                make_table(
                    content["tables"]["roadmap"],
                    [34 * mm, 34 * mm, 102 * mm],
                )
            )
        story.append(Spacer(1, 8))

    story.append(PageBreak())
    story.append(Paragraph("Источники", h1))
    for source in content["sources"]:
        story.append(Paragraph(source, bullet, bulletText="•"))
    story.append(Spacer(1, 8))
    story.append(
        Paragraph(
            f"Основание документа: репозиторные документы и research snapshot от {meta['basis_date']}.",
            small,
        )
    )

    OUTPUT_PDF.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story, onFirstPage=pdf_header_footer, onLaterPages=pdf_header_footer)


def main() -> None:
    content = load_content()
    build_docx(content)
    build_pdf(content)
    print(f"Generated {OUTPUT_DOCX}")
    print(f"Generated {OUTPUT_PDF}")


if __name__ == "__main__":
    main()
